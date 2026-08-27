"""
Document Ingestion & Semantic Section Parser.
Supports PDF, DOCX, TXT, and Markdown parsing with metadata extraction,
Document Passport generation, and BM25 relevance section scoring.
"""

from __future__ import annotations

import io
import math
import re
import uuid
from typing import Any

import pypdf
import docx

from app.agent.scoring import calculate_lexical_relevance


def parse_uploaded_file(
    file_bytes: bytes,
    filename: str,
    content_type: str | None = None,
) -> dict[str, Any]:
    """
    Parses a single uploaded document (PDF, DOCX, TXT, MD, CSV) into structured
    page-indexed text, computes statistics, and generates a Document Passport.
    """
    clean_name = filename.strip()
    ext = clean_name.lower().rsplit(".", 1)[-1] if "." in clean_name else "txt"
    file_size = len(file_bytes)

    pages: list[dict[str, Any]] = []
    full_text_parts: list[str] = []

    if ext == "pdf":
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            total_pages = len(reader.pages)
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                page_text = page_text.strip()
                if page_text:
                    pages.append({"page_num": idx + 1, "text": page_text})
                    full_text_parts.append(f"--- Page {idx + 1} ---\n{page_text}")
        except Exception as exc:
            # Fallback if corrupted or encrypted
            pages.append({"page_num": 1, "text": f"[Error reading PDF: {exc}]"})
            full_text_parts.append(f"[Error reading PDF: {exc}]")

    elif ext in ("docx", "doc"):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            para_texts: list[str] = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    para_texts.append(t)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        para_texts.append(" | ".join(row_cells))

            full_docx_text = "\n\n".join(para_texts)
            # Estimate pages based on ~350 words per page
            words = full_docx_text.split()
            page_size = 350
            total_pages = max(1, math.ceil(len(words) / page_size)) if words else 1

            for p_idx in range(total_pages):
                chunk_words = words[p_idx * page_size : (p_idx + 1) * page_size]
                chunk_text = " ".join(chunk_words)
                if chunk_text:
                    pages.append({"page_num": p_idx + 1, "text": chunk_text})
                    full_text_parts.append(f"--- Section / Page {p_idx + 1} ---\n{chunk_text}")
        except Exception as exc:
            pages.append({"page_num": 1, "text": f"[Error reading DOCX: {exc}]"})
            full_text_parts.append(f"[Error reading DOCX: {exc}]")

    else:
        # Plain text, Markdown, CSV, etc.
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="replace")

        text = text.strip()
        words = text.split()
        page_size = 400
        total_pages = max(1, math.ceil(len(words) / page_size)) if words else 1

        for p_idx in range(total_pages):
            chunk_words = words[p_idx * page_size : (p_idx + 1) * page_size]
            chunk_text = " ".join(chunk_words)
            if chunk_text:
                pages.append({"page_num": p_idx + 1, "text": chunk_text})
                full_text_parts.append(f"--- Section {p_idx + 1} ---\n{chunk_text}")

    combined_text = "\n\n".join(full_text_parts)
    total_words = len(combined_text.split())
    total_chars = len(combined_text)
    total_pages = max(1, len(pages))

    # Generate 1-2 sentence preview for the Document Passport
    first_page_text = pages[0]["text"] if pages else ""
    preview_clean = re.sub(r"\s+", " ", first_page_text).strip()
    preview = preview_clean[:280] + ("..." if len(preview_clean) > 280 else "")

    doc_id = f"doc_{uuid.uuid4().hex[:8]}"

    return {
        "id": doc_id,
        "filename": clean_name,
        "file_type": ext,
        "file_size": file_size,
        "page_count": total_pages,
        "word_count": total_words,
        "char_count": total_chars,
        "pages": pages,
        "full_text": combined_text,
        "preview": preview,
    }


def score_and_extract_relevant_sections(
    documents: list[dict[str, Any]],
    query: str,
    max_chars: int = 35000,
) -> str:
    """
    Ranks document pages across all attached files using BM25 relevance scoring
    against the research query, returning the highest-density grounded excerpts
    with explicit page markers.
    """
    if not documents:
        return ""

    # If all combined text fits easily within max_chars, return full formatted text
    total_len = sum(len(d.get("full_text", "")) for d in documents)
    if total_len <= max_chars:
        formatted_docs = []
        for d in documents:
            formatted_docs.append(
                f"### [DOCUMENT: {d['filename']}] (Format: {d.get('file_type', 'pdf').upper()}, "
                f"Pages: {d.get('page_count', 1)}, Words: {d.get('word_count', 0)})\n"
                f"{d.get('full_text', '')}"
            )
        return "\n\n".join(formatted_docs)

    # Document is large: break into page-level chunks and compute BM25 relevance
    scored_chunks: list[dict[str, Any]] = []

    for doc in documents:
        fname = doc.get("filename", "Uploaded_Document")
        pages = doc.get("pages", [])
        for p in pages:
            p_num = p.get("page_num", 1)
            p_text = p.get("text", "")
            if not p_text.strip():
                continue

            score, _ = calculate_lexical_relevance(p_text, query, query)
            scored_chunks.append({
                "filename": fname,
                "page_num": p_num,
                "text": p_text,
                "score": score,
            })

    # Sort chunks by relevance score descending
    scored_chunks.sort(key=lambda x: x["score"], reverse=True)

    selected_chunks: list[str] = []
    current_chars = 0

    for chunk in scored_chunks:
        chunk_header = f"[Doc: {chunk['filename']}, Page {chunk['page_num']}] (BM25 Relevance: {chunk['score']:.2f})"
        chunk_block = f"{chunk_header}\n{chunk['text']}"
        if current_chars + len(chunk_block) > max_chars:
            break
        selected_chunks.append(chunk_block)
        current_chars += len(chunk_block)

    if not selected_chunks and documents:
        # Fallback: take the first page of each document
        for doc in documents:
            p1 = doc.get("pages", [{}])[0].get("text", "")
            selected_chunks.append(f"[Doc: {doc['filename']}, Page 1]\n{p1[:max_chars // len(documents)]}")

    return "\n\n---\n\n".join(selected_chunks)
