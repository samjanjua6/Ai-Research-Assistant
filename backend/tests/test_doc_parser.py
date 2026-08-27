import io
import unittest
import docx
from pypdf import PdfWriter

from app.agent.doc_parser import parse_uploaded_file, score_and_extract_relevant_sections


class TestDocParser(unittest.TestCase):
    def test_txt_and_md_parsing(self):
        content = b"# Clinical Trial Phase 3\n\nResults show 85% efficacy with zero serious adverse events."
        doc = parse_uploaded_file(content, "trial_results.md")

        self.assertEqual(doc["filename"], "trial_results.md")
        self.assertEqual(doc["file_type"], "md")
        self.assertGreater(doc["word_count"], 5)
        self.assertIn("85% efficacy", doc["full_text"])
        self.assertTrue(len(doc["preview"]) > 10)

    def test_docx_parsing(self):
        # Create an in-memory docx file
        doc_obj = docx.Document()
        doc_obj.add_heading("Solid-State Battery Specifications", level=1)
        doc_obj.add_paragraph("Energy density reaches 450 Wh/kg with silicon-composite anodes.")
        table = doc_obj.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Cycle Life"
        table.cell(1, 1).text = "1200 cycles"

        buf = io.BytesIO()
        doc_obj.save(buf)
        buf.seek(0)

        parsed = parse_uploaded_file(buf.read(), "battery_specs.docx")
        self.assertEqual(parsed["filename"], "battery_specs.docx")
        self.assertEqual(parsed["file_type"], "docx")
        self.assertIn("450 Wh/kg", parsed["full_text"])
        self.assertIn("Cycle Life", parsed["full_text"])

    def test_pdf_parsing(self):
        # Create an in-memory blank/simple pdf
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)

        parsed = parse_uploaded_file(buf.read(), "sample_paper.pdf")
        self.assertEqual(parsed["filename"], "sample_paper.pdf")
        self.assertEqual(parsed["file_type"], "pdf")
        self.assertEqual(parsed["page_count"], 1)

    def test_bm25_semantic_section_scoring(self):
        docs = [
            {
                "id": "doc1",
                "filename": "quantum_chips.pdf",
                "file_type": "pdf",
                "page_count": 2,
                "word_count": 100,
                "pages": [
                    {"page_num": 1, "text": "Superconducting qubits require dilution refrigerators operating at 15 millikelvin."},
                    {"page_num": 2, "text": "Neutral atom systems use optical tweezers to trap rubidium atoms at room temperature."},
                ],
                "full_text": ("Superconducting qubits... " * 50) + ("Neutral atom systems... " * 50),
            }
        ]

        query = "How do neutral atom systems compare in operating temperature?"
        # Trigger chunk scoring by setting max_chars (300) < total_len (2500)
        extracted = score_and_extract_relevant_sections(docs, query, max_chars=300)
        self.assertIn("Neutral atom systems", extracted)
        self.assertIn("Page 2", extracted)


if __name__ == "__main__":
    unittest.main()
